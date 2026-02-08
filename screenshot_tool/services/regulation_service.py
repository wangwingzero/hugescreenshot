#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CAAC 规章查询服务

提供规章数据获取、搜索、PDF 下载功能。
数据来源：中国民用航空局官网 (https://www.caac.gov.cn)

Feature: caac-regulation-search
"""

import os
import re
from dataclasses import dataclass
from typing import Optional
from urllib.parse import urljoin, quote

import requests
from bs4 import BeautifulSoup
from PySide6.QtCore import QObject, QThread, Signal

# 调试日志
try:
    from screenshot_tool.core.async_logger import async_debug_log as _debug_log
except ImportError:
    def _debug_log(msg, tag="INFO"): print(f"[{tag}] {msg}")


# CAAC 官网配置
BASE_URL = "https://www.caac.gov.cn"
# WAS5 搜索 API - 实际的搜索接口
WAS5_SEARCH_URL = "https://www.caac.gov.cn/was5/web/search"

# 频道 ID
REGULATION_CHANNEL = "269689"  # 民航规章频道
NORMATIVE_CHANNEL = "238066"   # 规范性文件频道

# 分类 ID (fl 参数)
REGULATION_FL = "13"   # 民航规章分类
NORMATIVE_FL = "14"    # 规范性文件分类

DEFAULT_HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
    'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
    'Accept-Encoding': 'gzip, deflate, br',
    'Connection': 'keep-alive',
    'Referer': 'https://www.caac.gov.cn/XXGK/XXGK/',
}


@dataclass
class RegulationDocument:
    """规章文档数据模型
    
    Attributes:
        title: 文档标题
        url: 详情页 URL
        validity: 有效性状态 ("有效", "失效", "废止")
        doc_number: 文号 (如 "CCAR-121-R8", "AC-121-FS-139")
        office_unit: 发布单位 (如 "政策法规司", "飞行标准司")
        doc_type: 文档类型 ("regulation" 规章, "normative" 规范性文件)
        sign_date: 签发日期 (规范性文件专用)
        publish_date: 发布日期
        file_number: 字号 (规范性文件专用，如 "民航规〔2025〕5 号")
        pdf_url: PDF 附件下载链接 (如果有)
    """
    title: str
    url: str
    validity: str
    doc_number: str
    office_unit: str
    doc_type: str  # "regulation" or "normative"
    sign_date: str = ""
    publish_date: str = ""
    file_number: str = ""
    pdf_url: str = ""  # PDF 附件链接
    
    def to_dict(self) -> dict:
        """转换为字典"""
        return {
            "title": self.title,
            "url": self.url,
            "validity": self.validity,
            "doc_number": self.doc_number,
            "office_unit": self.office_unit,
            "doc_type": self.doc_type,
            "sign_date": self.sign_date,
            "publish_date": self.publish_date,
            "file_number": self.file_number,
        }
    
    @classmethod
    def from_dict(cls, data: dict) -> "RegulationDocument":
        """从字典创建"""
        return cls(
            title=data.get("title", ""),
            url=data.get("url", ""),
            validity=data.get("validity", ""),
            doc_number=data.get("doc_number", ""),
            office_unit=data.get("office_unit", ""),
            doc_type=data.get("doc_type", "regulation"),
            sign_date=data.get("sign_date", ""),
            publish_date=data.get("publish_date", ""),
            file_number=data.get("file_number", ""),
        )


def generate_filename(document: RegulationDocument) -> str:
    """生成 PDF 文件名
    
    命名规则：
    - 规范性文件（有效）: {文号}{名称}.pdf
      例: AC-21-AA-2008-15运营人航空器适航检查单.pdf
    - 规范性文件（失效/废止）: 失效!{文号}{名称}.pdf
      例: 失效!AC-00-AA-2009-02R1适航培训课程目录.pdf
    - CCAR规章（有效）: {CCAR部号}{名称}.pdf
      例: CCAR-25-R4运输类飞机适航标准.pdf
    - CCAR规章（失效/废止）: 失效!{CCAR部号}{名称}.pdf
      例: 失效!CCAR-14-R1民用航空行政处罚实施办法.pdf
    
    Args:
        document: 文档对象
        
    Returns:
        生成的文件名
        
    Feature: caac-regulation-download-enhancement
    Property 3: Normative Document Filename Generation
    Property 4: CCAR Regulation Filename Generation
    Validates: Requirements 4.1-4.5, 5.1-5.5
    """
    # 清理特殊字符
    def sanitize(text: str) -> str:
        """替换文件名中的非法字符"""
        return re.sub(r'[<>:"/\\|?*]', '_', text)
    
    parts = []
    
    # 有效性前缀 - 失效和废止统一使用 "失效!" 前缀
    validity = document.validity.strip()
    if validity in ("失效", "废止"):
        parts.append("失效!")
    
    # 文号（规范性文件用文号，CCAR规章用CCAR部号）
    doc_number = sanitize(document.doc_number.strip())
    if doc_number:
        parts.append(doc_number)
    
    # 标题
    title = sanitize(document.title.strip())
    parts.append(title)
    
    # 组合文件名
    filename = "".join(parts) + ".pdf"
    
    # 限制文件名长度
    if len(filename) > 200:
        filename = filename[:197] + "....pdf"
    
    return filename


def get_save_path(document: RegulationDocument, base_path: str = "") -> str:
    """获取文档保存路径
    
    直接保存到 base_path 目录，不创建子目录。
    
    Args:
        document: 文档对象
        base_path: 基础保存路径，默认为 ~/Documents/CAAC_PDF/
        
    Returns:
        完整的保存路径（包含文件名）
        
    Feature: caac-regulation-search
    Validates: Requirements 6.1, 6.3, 6.4
    """
    # 默认保存路径
    if not base_path:
        base_path = os.path.join(os.path.expanduser("~"), "Documents", "CAAC_PDF")
    
    # 创建目录
    os.makedirs(base_path, exist_ok=True)
    
    # 生成文件名
    filename = generate_filename(document)
    
    return os.path.join(base_path, filename)


class RegulationSearchWorker(QThread):
    """规章搜索工作线程
    
    直接调用 CAAC 官网搜索功能。
    由于 CAAC 网站使用 JavaScript 反爬虫保护，优先使用浏览器模式获取内容。
    
    Feature: caac-regulation-search
    """
    
    # 信号
    finished = Signal(list)  # 搜索完成，参数为文档列表
    error = Signal(str)  # 搜索错误
    progress = Signal(int, int, str)  # 进度 (current, total, message)
    
    def __init__(
        self,
        keyword: str,
        doc_type: str = "all",
        validity: str = "all",
        start_date: str = "",
        end_date: str = "",
        session: Optional[requests.Session] = None
    ):
        """初始化
        
        Args:
            keyword: 搜索关键词
            doc_type: 文档类型 ("all", "regulation", "normative")
            validity: 有效性 ("all", "valid", "invalid")
            start_date: 起始日期 (YYYY-MM-DD)，空字符串表示不限制
            end_date: 结束日期 (YYYY-MM-DD)，空字符串表示不限制
            session: HTTP 会话
        """
        super().__init__()
        self._keyword = keyword
        self._doc_type = doc_type
        self._validity = validity
        self._start_date = start_date
        self._end_date = end_date
        self._session = session or requests.Session()
        self._session.headers.update(DEFAULT_HEADERS)
        self._should_stop = False
        self._browser_fetcher = None
        self._browser_fetcher_checked = False  # 避免重复检查
    
    def _get_browser_fetcher(self):
        """获取浏览器获取器（懒加载，线程安全）
        
        Returns:
            BrowserFetcher 实例或 None
        """
        # 只检查一次，避免重复导入和检查
        if self._browser_fetcher_checked:
            return self._browser_fetcher
        
        self._browser_fetcher_checked = True
        
        try:
            from screenshot_tool.services.browser_fetcher import BrowserFetcher
            if BrowserFetcher.is_available():
                self._browser_fetcher = BrowserFetcher(timeout=30)
                _debug_log("浏览器模式可用", "REGULATION")
            else:
                _debug_log("浏览器模式不可用（未找到浏览器或 patchright）", "REGULATION")
        except ImportError as e:
            _debug_log(f"browser_fetcher 模块导入失败: {e}", "REGULATION")
        except Exception as e:
            _debug_log(f"初始化浏览器获取器失败: {e}", "REGULATION")
        
        return self._browser_fetcher
    
    def stop(self):
        """请求停止"""
        self._should_stop = True
    
    def run(self):
        """执行搜索"""
        try:
            documents = []
            
            # 根据文档类型决定搜索哪些页面
            if self._doc_type in ("all", "regulation"):
                self.progress.emit(0, 100, "正在搜索 CCAR 规章...")
                regulations = self._search_regulations()
                documents.extend(regulations)
                _debug_log(f"CCAR 规章搜索完成，找到 {len(regulations)} 条", "REGULATION")
            
            if self._should_stop:
                return
            
            if self._doc_type in ("all", "normative"):
                self.progress.emit(50, 100, "正在搜索规范性文件...")
                normatives = self._search_normatives()
                documents.extend(normatives)
                _debug_log(f"规范性文件搜索完成，找到 {len(normatives)} 条", "REGULATION")
            
            if self._should_stop:
                return
            
            # 根据有效性筛选
            if self._validity == "valid":
                documents = [d for d in documents if d.validity == "有效"]
            elif self._validity == "invalid":
                documents = [d for d in documents if d.validity in ("失效", "废止")]
            
            self.progress.emit(100, 100, f"搜索完成，找到 {len(documents)} 个结果")
            self.finished.emit(documents)
            
        except Exception as e:
            _debug_log(f"搜索失败: {e}", "REGULATION")
            self.error.emit(f"搜索失败: {str(e)}")
    
    def _fetch_page_content(self, url: str) -> str:
        """获取页面内容
        
        优先使用浏览器模式（绕过 JS 反爬虫），失败时回退到 requests。
        
        Args:
            url: 页面 URL
            
        Returns:
            HTML 内容，失败时返回空字符串
        """
        _debug_log(f"获取页面: {url}", "REGULATION")
        
        # 优先使用浏览器模式
        browser_fetcher = self._get_browser_fetcher()
        if browser_fetcher:
            try:
                result = browser_fetcher.fetch(url, use_cookies=False)
                if result.success and result.html:
                    _debug_log(f"浏览器模式获取成功: {len(result.html)} 字符", "REGULATION")
                    return result.html
                else:
                    error_msg = result.error if hasattr(result, 'error') else "未知错误"
                    _debug_log(f"浏览器模式获取失败: {error_msg}", "REGULATION")
            except Exception as e:
                _debug_log(f"浏览器模式异常: {e}", "REGULATION")
        
        # 回退到 requests（可能无法获取 JS 渲染的内容）
        _debug_log("回退到 requests 模式", "REGULATION")
        try:
            response = self._session.get(url, timeout=30)
            response.encoding = "utf-8"
            if response.status_code == 200:
                return response.text
        except Exception as e:
            _debug_log(f"requests 获取失败: {e}", "REGULATION")
        
        return ""
    
    def _build_date_params(self) -> str:
        """构建日期参数字符串
        
        WAS5 搜索系统使用 fwrq1 和 fwrq2 参数进行日期筛选。
        日期格式：YYYY-MM-DD
        
        Returns:
            URL 参数字符串，如 "&fwrq1=2025-01-01&fwrq2=2025-01-09"
        """
        params = ""
        if self._start_date:
            params += f"&fwrq1={self._start_date}"
        if self._end_date:
            params += f"&fwrq2={self._end_date}"
        return params
    
    def _search_regulations(self) -> list:
        """搜索 CCAR 规章"""
        documents = []
        
        try:
            # 构建 WAS5 搜索 URL
            # 规章使用 channelid=269689, fl=13
            date_params = self._build_date_params()
            
            if self._keyword:
                # 使用 sw 参数搜索（全文搜索，包含标题）
                search_url = f"{WAS5_SEARCH_URL}?channelid={REGULATION_CHANNEL}&sw={quote(self._keyword)}&perpage=100&orderby=-fabuDate&fl={REGULATION_FL}{date_params}"
            else:
                # 无关键词时，获取所有规章
                search_url = f"{WAS5_SEARCH_URL}?channelid={REGULATION_CHANNEL}&perpage=100&orderby=-fabuDate&fl={REGULATION_FL}{date_params}"
            
            _debug_log(f"规章搜索 URL: {search_url}", "REGULATION")
            html_content = self._fetch_page_content(search_url)
            
            if html_content:
                documents = self._parse_regulation_page(html_content)
                _debug_log(f"解析规章页面，找到 {len(documents)} 条记录", "REGULATION")
            else:
                _debug_log("获取规章页面失败，内容为空", "REGULATION")
                
        except Exception as e:
            _debug_log(f"搜索规章失败: {e}", "REGULATION")
        
        return documents
    
    def _search_normatives(self) -> list:
        """搜索规范性文件"""
        documents = []
        
        try:
            # 构建 WAS5 搜索 URL
            # 规范性文件使用 channelid=238066, fl=14
            date_params = self._build_date_params()
            
            if self._keyword:
                # 使用 sw 参数搜索（全文搜索，包含标题）
                search_url = f"{WAS5_SEARCH_URL}?channelid={NORMATIVE_CHANNEL}&sw={quote(self._keyword)}&perpage=100&orderby=-fabuDate&fl={NORMATIVE_FL}{date_params}"
            else:
                # 无关键词时，获取所有规范性文件
                search_url = f"{WAS5_SEARCH_URL}?channelid={NORMATIVE_CHANNEL}&perpage=100&orderby=-fabuDate&fl={NORMATIVE_FL}{date_params}"
            
            _debug_log(f"规范性文件搜索 URL: {search_url}", "REGULATION")
            html_content = self._fetch_page_content(search_url)
            
            if html_content:
                documents = self._parse_normative_page(html_content)
                _debug_log(f"解析规范性文件页面，找到 {len(documents)} 条记录", "REGULATION")
            else:
                _debug_log("获取规范性文件页面失败，内容为空", "REGULATION")
                
        except Exception as e:
            _debug_log(f"搜索规范性文件失败: {e}", "REGULATION")
        
        return documents
    
    def _parse_regulation_page(self, html_content: str) -> list:
        """解析规章搜索结果页面
        
        CAAC 网站 CCAR 规章表格结构：
        - 表格 class="t_table"
        - 列: 序号 | 名称(class="t_l") | CCAR部号 | 有效性
        - 发文日期从 URL 中提取（格式如 202512/t20251217_229435.html）
        """
        documents = []
        
        def normalize_date(date_str: str) -> str:
            """将中文日期格式转换为 YYYY-MM-DD 格式
            
            例如: "2025年12月26日" -> "2025-12-26"
                  "2025-11-27" -> "2025-11-27"
            """
            if not date_str:
                return ""
            # 匹配中文日期格式
            match = re.match(r'(\d{4})年(\d{1,2})月(\d{1,2})日', date_str)
            if match:
                year, month, day = match.groups()
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
            # 如果已经是 YYYY-MM-DD 格式，直接返回
            if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                return date_str
            return date_str
        
        def extract_date_from_url(url: str) -> str:
            """从 URL 中提取发文日期
            
            URL 格式如: http://www.caac.gov.cn/XXGK/XXGK/MHGZ/202512/t20251217_229435.html
            其中 t20251217 表示 2025年12月17日
            
            Returns:
                YYYY-MM-DD 格式的日期，提取失败返回空字符串
            """
            # 匹配 URL 中的日期部分: /tYYYYMMDD_
            match = re.search(r'/t(\d{4})(\d{2})(\d{2})_', url)
            if match:
                year, month, day = match.groups()
                return f"{year}-{month}-{day}"
            return ""
        
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            
            # 查找主数据表格 (class="t_table")
            table = soup.find("table", class_="t_table")
            _debug_log(f"查找 t_table 表格: {'找到' if table else '未找到'}", "REGULATION")
            
            if not table:
                # 回退：查找任意表格
                tables = soup.find_all("table")
                _debug_log(f"回退查找所有表格，找到 {len(tables)} 个", "REGULATION")
                table = tables[0] if tables else None
            
            if not table:
                _debug_log("未找到任何表格", "REGULATION")
                return documents
            
            # 查找 tbody 中的数据行
            tbody = table.find("tbody")
            if tbody:
                rows = tbody.find_all("tr")
                _debug_log(f"在 tbody 中找到 {len(rows)} 行", "REGULATION")
            else:
                rows = table.find_all("tr")[1:]  # 跳过表头
                _debug_log(f"直接在 table 中找到 {len(rows)} 行（跳过表头）", "REGULATION")
            
            for row in rows:
                if self._should_stop:
                    break
                
                cells = row.find_all("td")
                if len(cells) < 4:
                    continue
                
                try:
                    # CCAR 规章页面结构: 序号(0) | 名称(1) | CCAR部号(2) | 有效性(3)
                    # 名称单元格 class="t_l"
                    title_cell = row.find("td", class_="t_l")
                    if not title_cell and len(cells) > 1:
                        title_cell = cells[1]
                    if not title_cell:
                        continue
                    
                    link = title_cell.find("a", href=True)
                    if not link:
                        continue
                    
                    title = link.get_text(strip=True)
                    href = link.get("href", "")
                    full_url = urljoin(BASE_URL, href)
                    
                    # 提取 CCAR 部号 - 第 3 列 (index 2)
                    doc_number = ""
                    if len(cells) > 2:
                        doc_number = cells[2].get_text(strip=True)
                    
                    # 提取有效性 - 第 4 列 (index 3)
                    validity = ""
                    if len(cells) > 3:
                        validity = cells[3].get_text(strip=True)
                    
                    # 从 URL 中提取发文日期
                    publish_date = extract_date_from_url(full_url)
                    
                    # 从详情区域提取办文单位（如果有的话）
                    office_unit = ""
                    detail_div = title_cell.find("div", class_="t_l_content")
                    if detail_div:
                        # 查找所有 li 元素
                        for li in detail_div.find_all("li"):
                            li_text = li.get_text(strip=True)
                            if "办文单位：" in li_text:
                                office_unit = li_text.replace("办文单位：", "").strip()
                            elif "发文日期：" in li_text or "发文日期:" in li_text:
                                # 如果详情区域有发文日期，优先使用
                                date_text = re.sub(r"发文日期[：:]", "", li_text).strip()
                                publish_date = normalize_date(date_text) or publish_date
                            elif "有效性" in li_text and not validity:
                                # 从详情区域获取有效性（备用，处理各种空格情况）
                                validity = re.sub(r"有\s*效\s*性\s*[：:]", "", li_text).strip()
                    
                    _debug_log(f"解析规章: {title}, 部号: {doc_number}, 有效性: {validity}, 发文日期: {publish_date}", "REGULATION")
                    
                    if title:
                        documents.append(RegulationDocument(
                            title=title,
                            url=full_url,
                            validity=validity,
                            doc_number=doc_number,
                            office_unit=office_unit,
                            doc_type="regulation",
                            publish_date=publish_date,
                        ))
                except Exception as e:
                    _debug_log(f"解析行失败: {e}", "REGULATION")
                    continue
                    
        except Exception as e:
            _debug_log(f"解析规章页面失败: {e}", "REGULATION")
        
        return documents
    
    def _parse_normative_page(self, html_content: str) -> list:
        """解析规范性文件搜索结果页面
        
        CAAC 网站表格结构：
        - 表格 class="t_table"
        - 列: 序号 | 名称(tdMC) | 成文日期(tdRQ) | 发文日期(tdRQ) | 文号(strFL) | 有效性(strGF)
        - 详情区域 div.t_l_content 包含办文单位等信息
        """
        documents = []
        
        def normalize_date(date_str: str) -> str:
            """将中文日期格式转换为 YYYY-MM-DD 格式
            
            例如: "2025年12月26日" -> "2025-12-26"
            """
            if not date_str:
                return ""
            # 匹配中文日期格式
            match = re.match(r'(\d{4})年(\d{1,2})月(\d{1,2})日', date_str)
            if match:
                year, month, day = match.groups()
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
            # 如果已经是 YYYY-MM-DD 格式，直接返回
            if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                return date_str
            return date_str
        
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            
            # 查找主数据表格 (class="t_table")
            table = soup.find("table", class_="t_table")
            _debug_log(f"查找 t_table 表格: {'找到' if table else '未找到'}", "REGULATION")
            
            if not table:
                # 回退：查找任意表格
                tables = soup.find_all("table")
                _debug_log(f"回退查找所有表格，找到 {len(tables)} 个", "REGULATION")
                table = tables[0] if tables else None
            
            if not table:
                _debug_log("未找到任何表格", "REGULATION")
                return documents
            
            # 查找 tbody 中的数据行（跳过 thead）
            tbody = table.find("tbody")
            if tbody:
                rows = tbody.find_all("tr")
                _debug_log(f"在 tbody 中找到 {len(rows)} 行", "REGULATION")
            else:
                rows = table.find_all("tr")[1:]  # 跳过表头
                _debug_log(f"直接在 table 中找到 {len(rows)} 行（跳过表头）", "REGULATION")
            
            for row in rows:
                if self._should_stop:
                    break
                
                cells = row.find_all("td")
                if len(cells) < 4:
                    continue
                
                try:
                    # 提取链接和标题 - 在 tdMC 单元格中
                    title_cell = row.find("td", class_="tdMC")
                    if not title_cell and len(cells) > 1:
                        title_cell = cells[1]
                    if not title_cell:
                        continue
                    
                    link = title_cell.find("a", href=True)
                    if not link:
                        continue
                    
                    title = link.get_text(strip=True)
                    href = link.get("href", "")
                    full_url = urljoin(BASE_URL, href)
                    
                    # 提取文号 - 在 strFL 单元格中
                    doc_number = ""
                    doc_number_cell = row.find("td", class_="strFL")
                    if doc_number_cell:
                        doc_number = doc_number_cell.get_text(strip=True)
                    
                    # 提取有效性 - 在 strGF 单元格中
                    validity = ""
                    validity_cell = row.find("td", class_="strGF")
                    if validity_cell:
                        validity = validity_cell.get_text(strip=True)
                    
                    # 提取日期 - 在 tdRQ 单元格中，转换为 YYYY-MM-DD 格式
                    sign_date = ""
                    publish_date = ""
                    date_cells = row.find_all("td", class_="tdRQ")
                    if len(date_cells) >= 1:
                        sign_date = normalize_date(date_cells[0].get_text(strip=True))
                    if len(date_cells) >= 2:
                        publish_date = normalize_date(date_cells[1].get_text(strip=True))
                    
                    # 从详情区域提取办文单位
                    office_unit = ""
                    detail_div = title_cell.find("div", class_="t_l_content")
                    if detail_div:
                        unit_li = detail_div.find("li", class_="t_l_content_left")
                        if unit_li:
                            unit_text = unit_li.get_text(strip=True)
                            if "办文单位：" in unit_text:
                                office_unit = unit_text.replace("办文单位：", "").strip()
                    
                    if title:
                        documents.append(RegulationDocument(
                            title=title,
                            url=full_url,
                            validity=validity,
                            doc_number=doc_number,
                            office_unit=office_unit,
                            doc_type="normative",
                            sign_date=sign_date,
                            publish_date=publish_date,
                        ))
                except Exception as e:
                    _debug_log(f"解析行失败: {e}", "REGULATION")
                    continue
                    
        except Exception as e:
            _debug_log(f"解析规范性文件页面失败: {e}", "REGULATION")
        
        return documents
        
        return documents


class RegulationService(QObject):
    """规章服务 - 搜索和 PDF 下载
    
    Feature: caac-regulation-search
    Feature: caac-regulation-download-enhancement
    """
    
    # 信号
    searchFinished = Signal(list)  # 搜索完成，参数为文档列表
    searchError = Signal(str)  # 搜索错误，参数为错误信息
    searchProgress = Signal(int, int, str)  # 搜索进度 (current, total, message)
    downloadProgress = Signal(int, int, str)  # 下载进度 (current, total, message)
    downloadComplete = Signal(str)  # 下载完成，参数为文件路径
    downloadError = Signal(str)  # 下载错误，参数为错误信息
    
    def __init__(self, save_path: str = ""):
        """初始化服务
        
        Args:
            save_path: PDF 保存路径，默认为 ~/Documents/CAAC_PDF/
        """
        super().__init__()
        
        self._save_path = save_path or os.path.join(
            os.path.expanduser("~"), "Documents", "CAAC_PDF"
        )
        
        # HTTP 会话
        self._session = requests.Session()
        self._session.headers.update(DEFAULT_HEADERS)
        
        # 工作线程 - 使用列表保存旧线程引用，防止被垃圾回收
        self._search_worker: Optional[RegulationSearchWorker] = None
        self._download_worker: Optional["PDFDownloadWorker"] = None
        self._old_workers: list = []  # 保存旧线程引用，防止运行中被销毁
    
    @property
    def save_path(self) -> str:
        """获取保存路径"""
        return self._save_path
    
    @save_path.setter
    def save_path(self, value: str):
        """设置保存路径"""
        self._save_path = value
    
    def set_save_path(self, path: str):
        """设置保存路径
        
        Args:
            path: 新的保存路径
        """
        self._save_path = path
    
    def search(
        self,
        keyword: str = "",
        doc_type: str = "all",
        validity: str = "all",
        start_date: str = "",
        end_date: str = "",
    ) -> None:
        """搜索规章（异步）
        
        直接调用 CAAC 官网搜索功能。
        
        Args:
            keyword: 搜索关键词
            doc_type: 文档类型 ("all", "regulation", "normative")
            validity: 有效性 ("all", "valid", "invalid")
            start_date: 起始日期 (YYYY-MM-DD)，空字符串表示不限制
            end_date: 结束日期 (YYYY-MM-DD)，空字符串表示不限制
        """
        # 清理已完成的旧线程
        self._old_workers = [w for w in self._old_workers if w.isRunning()]
        
        # 限制旧线程列表大小，防止内存泄漏
        if len(self._old_workers) > 10:
            # 强制等待最旧的线程完成
            oldest = self._old_workers.pop(0)
            oldest.stop()
            oldest.wait(500)
        
        # 停止之前的搜索
        if self._search_worker and self._search_worker.isRunning():
            self._search_worker.stop()
            # 将旧线程移到列表中保存引用，防止被垃圾回收导致崩溃
            self._old_workers.append(self._search_worker)
            # 断开信号连接，避免旧线程的信号干扰
            try:
                self._search_worker.finished.disconnect(self._on_search_finished)
                self._search_worker.error.disconnect(self._on_search_error)
                self._search_worker.progress.disconnect(self._on_search_progress)
            except (RuntimeError, TypeError):
                # 信号可能未连接或已断开
                pass
        
        # 创建新的搜索线程
        self._search_worker = RegulationSearchWorker(
            keyword=keyword,
            doc_type=doc_type,
            validity=validity,
            start_date=start_date,
            end_date=end_date,
            session=self._session,
        )
        self._search_worker.finished.connect(self._on_search_finished)
        self._search_worker.error.connect(self._on_search_error)
        self._search_worker.progress.connect(self._on_search_progress)
        self._search_worker.start()
    
    def _on_search_finished(self, documents: list):
        """搜索完成"""
        self.searchFinished.emit(documents)
    
    def _on_search_error(self, error_msg: str):
        """搜索错误"""
        self.searchError.emit(error_msg)
    
    def _on_search_progress(self, current: int, total: int, message: str):
        """搜索进度"""
        self.searchProgress.emit(current, total, message)
    
    def download_pdf(self, document: RegulationDocument) -> None:
        """下载文档（异步）
        
        下载策略：
        1. 先检查详情页是否有 PDF 附件
        2. 如果有 PDF，直接下载 PDF
        3. 如果没有 PDF，提取内容生成 DOCX
        
        Args:
            document: 要下载的文档
            
        Feature: caac-regulation-download-enhancement
        Validates: Requirements 2.1, 3.1
        """
        # 停止之前的下载
        if self._download_worker and self._download_worker.isRunning():
            self._download_worker.stop()
            self._download_worker.wait(1000)
        
        # 获取保存路径
        save_path = get_save_path(document, self._save_path)
        
        # 统一使用智能下载 Worker（先检查 PDF，没有则生成 DOCX）
        self._download_worker = SmartDownloadWorker(
            self._session, document, save_path
        )
        
        self._download_worker.finished.connect(self._on_download_finished)
        self._download_worker.error.connect(self._on_download_error)
        self._download_worker.progress.connect(self._on_download_progress)
        self._download_worker.start()
    
    def _on_download_finished(self, file_path: str):
        """下载完成"""
        self.downloadComplete.emit(file_path)
    
    def _on_download_error(self, error_msg: str):
        """下载错误"""
        self.downloadError.emit(error_msg)
    
    def _on_download_progress(self, current: int, total: int, message: str = ""):
        """下载进度
        
        Args:
            current: 当前进度
            total: 总进度
            message: 进度消息（可选）
            
        Feature: caac-regulation-download-enhancement
        Validates: Requirements 6.1, 6.2
        """
        self.downloadProgress.emit(current, total, message)
    
    def cleanup(self):
        """清理资源"""
        # 停止当前搜索线程
        if self._search_worker and self._search_worker.isRunning():
            self._search_worker.stop()
            self._search_worker.wait(2000)
        
        # 停止当前下载线程
        if self._download_worker and self._download_worker.isRunning():
            self._download_worker.stop()
            self._download_worker.wait(2000)
        
        # 等待所有旧线程完成
        for worker in self._old_workers:
            if worker.isRunning():
                worker.stop()
                worker.wait(1000)
        self._old_workers.clear()


class PDFDownloadWorker(QThread):
    """PDF 下载工作线程
    
    Feature: caac-regulation-search
    Validates: Requirements 5.2, 5.6, 5.7
    """
    
    # 信号
    finished = Signal(str)  # 下载完成，参数为文件路径
    error = Signal(str)  # 下载错误
    progress = Signal(int, int)  # 进度 (downloaded, total)
    
    def __init__(
        self,
        session: requests.Session,
        document: RegulationDocument,
        save_path: str,
    ):
        """初始化
        
        Args:
            session: HTTP 会话
            document: 要下载的文档
            save_path: 保存路径
        """
        super().__init__()
        self._session = session
        self._document = document
        self._save_path = save_path
        self._should_stop = False
    
    def stop(self):
        """请求停止"""
        self._should_stop = True
    
    def run(self):
        """执行下载"""
        try:
            # 访问详情页
            response = self._session.get(self._document.url, timeout=30)
            response.encoding = "utf-8"
            
            if response.status_code != 200:
                self.error.emit(f"访问详情页失败: HTTP {response.status_code}")
                return
            
            soup = BeautifulSoup(response.text, "html.parser")
            
            # 查找 PDF 链接
            pdf_link = self._find_pdf_link(soup)
            
            if pdf_link:
                # 下载 PDF
                success = self._download_file(pdf_link, self._save_path)
                if success:
                    self.finished.emit(self._save_path)
                else:
                    # PDF 下载失败，尝试提取正文
                    txt_path = self._save_path.replace(".pdf", ".txt")
                    if self._extract_text_fallback(soup, txt_path):
                        self.finished.emit(txt_path)
                    else:
                        self.error.emit("PDF 下载失败，正文提取也失败")
            else:
                # 没有 PDF 链接，提取正文
                txt_path = self._save_path.replace(".pdf", ".txt")
                if self._extract_text_fallback(soup, txt_path):
                    self.finished.emit(txt_path)
                else:
                    self.error.emit("未找到 PDF 链接，正文提取也失败")
                    
        except Exception as e:
            self.error.emit(f"下载失败: {str(e)}")
    
    def _find_pdf_link(self, soup: BeautifulSoup) -> Optional[str]:
        """从详情页查找 PDF 链接
        
        参考法规文件项目的实现，使用多种模式查找 PDF/DOC 链接。
        
        Args:
            soup: BeautifulSoup 对象
            
        Returns:
            PDF/DOC 链接或 None
        """
        pdf_link = None
        
        # 模式1: 查找附件区域的 PDF 链接
        attachment_elements = soup.find_all(string=re.compile(r'附件[：:]?', re.I))
        for attachment_text in attachment_elements:
            attachment_parent = attachment_text.parent
            if attachment_parent:
                parent_container = attachment_parent.parent if attachment_parent.parent else attachment_parent
                pdf_links = parent_container.find_all('a', href=re.compile(r'\.pdf$', re.I))
                if pdf_links:
                    pdf_link = pdf_links[0].get('href')
                    break
        
        # 模式2: 直接查找所有 PDF 链接
        if not pdf_link:
            pdf_links = soup.find_all('a', href=re.compile(r'\.pdf$', re.I))
            if pdf_links:
                for link in pdf_links:
                    href = link.get('href', '')
                    if href and ('/XXGK/' in href or href.startswith('http')):
                        pdf_link = href
                        break
                if not pdf_link:
                    pdf_link = pdf_links[0].get('href')
        
        # 模式3: 查找 DOC 文件
        if not pdf_link:
            doc_links = soup.find_all('a', href=re.compile(r'\.(doc|docx)$', re.I))
            if doc_links:
                for link in doc_links:
                    href = link.get('href', '')
                    if href and ('/XXGK/' in href or href.startswith('http')):
                        pdf_link = href
                        break
                if not pdf_link:
                    pdf_link = doc_links[0].get('href')
        
        # 模式4: 包含下载关键词的链接
        if not pdf_link:
            download_links = soup.find_all('a', string=re.compile(r'(下载|PDF|pdf|DOC|doc|附件)', re.I))
            for link in download_links:
                href = link.get('href', '')
                if href and (
                    '.pdf' in href.lower()
                    or '.doc' in href.lower()
                    or 'download' in href.lower()
                    or 'attachment' in href.lower()
                ):
                    pdf_link = href
                    break
        
        # 模式5: onclick 事件中的 PDF 链接
        if not pdf_link:
            onclick_links = soup.find_all('a', onclick=re.compile(r'(\.pdf|download|attachment)', re.I))
            if onclick_links:
                onclick = onclick_links[0].get('onclick', '')
                pdf_match = re.search(r"['\"]([^'\"]*(?:\.pdf|download|attachment)[^'\"]*)['\"]", onclick, re.I)
                if pdf_match:
                    pdf_link = pdf_match.group(1)
        
        # 模式6: 所有包含文件下载相关的链接
        if not pdf_link:
            all_links = soup.find_all('a', href=True)
            for link in all_links:
                href = link.get('href', '')
                text = link.get_text(strip=True)
                if (
                    href
                    and (
                        'download' in href.lower()
                        or 'attachment' in href.lower()
                        or 'file' in href.lower()
                        or '.pdf' in href.lower()
                    )
                ) or (text and ('下载' in text or '附件' in text or 'PDF' in text.upper())):
                    pdf_link = href
                    break
        
        # 构建完整 URL
        if pdf_link:
            return self._build_full_url(pdf_link)
        
        return None
    
    def _build_full_url(self, link: str) -> str:
        """构建完整的 URL
        
        使用文档详情页 URL 的目录作为基准来解析相对路径。
        参考法规文件项目的 _build_pdf_url 实现。
        
        Args:
            link: 相对或绝对链接
            
        Returns:
            完整的 URL
        """
        if link.startswith('http'):
            return link
        
        # 处理相对路径前缀
        if link.startswith('./'):
            link = link[2:]
        elif link.startswith('../'):
            # 处理多级 ../ 
            doc_url = self._document.url
            while link.startswith('../'):
                link = link[3:]
                doc_url = '/'.join(doc_url.split('/')[:-1])
            doc_url_dir = '/'.join(doc_url.split('/')[:-1])
            return f"{doc_url_dir}/{link}"
        
        # 使用文档详情页 URL 的目录作为基准
        doc_url_dir = '/'.join(self._document.url.split('/')[:-1])
        return f"{doc_url_dir}/{link}"
    
    def _download_file(self, url: str, save_path: str) -> bool:
        """下载文件
        
        Args:
            url: 文件 URL
            save_path: 保存路径
            
        Returns:
            是否成功
        """
        try:
            # 根据 URL 调整文件扩展名
            url_lower = url.lower()
            if url_lower.endswith('.doc') or url_lower.endswith('.docx'):
                if save_path.endswith('.pdf'):
                    ext = '.docx' if url_lower.endswith('.docx') else '.doc'
                    save_path = save_path.replace('.pdf', ext)
                    # 更新保存路径供后续使用
                    self._save_path = save_path
            
            response = self._session.get(url, stream=True, timeout=60)
            
            if response.status_code != 200:
                return False
            
            # 获取文件大小
            total_size = int(response.headers.get("content-length", 0))
            downloaded = 0
            
            # 确保目录存在
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            # 写入文件
            try:
                with open(save_path, "wb") as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if self._should_stop:
                            # 清理未完成的文件
                            f.close()
                            if os.path.exists(save_path):
                                os.remove(save_path)
                            return False
                        
                        if chunk:
                            f.write(chunk)
                            downloaded += len(chunk)
                            
                            if total_size > 0:
                                self.progress.emit(downloaded, total_size)
            except Exception:
                # 清理未完成的文件
                if os.path.exists(save_path):
                    try:
                        os.remove(save_path)
                    except OSError:
                        pass
                raise
            
            # 验证文件大小
            file_size = os.path.getsize(save_path)
            if file_size < 1024:  # 小于 1KB 可能是错误页面
                os.remove(save_path)
                return False
            
            return True
            
        except Exception as e:
            print(f"下载文件失败: {e}")
            return False
    
    def _extract_text_fallback(self, soup: BeautifulSoup, save_path: str) -> bool:
        """提取正文作为备选方案
        
        Args:
            soup: BeautifulSoup 对象
            save_path: 保存路径
            
        Returns:
            是否成功
        """
        try:
            # 查找正文内容
            content_div = soup.find("div", class_="TRS_Editor")
            if not content_div:
                content_div = soup.find("div", class_="content")
            if not content_div:
                content_div = soup.find("div", id="content")
            
            if content_div:
                text = content_div.get_text(separator="\n", strip=True)
                
                if text:
                    # 确保目录存在
                    os.makedirs(os.path.dirname(save_path), exist_ok=True)
                    
                    with open(save_path, "w", encoding="utf-8") as f:
                        f.write(f"标题: {self._document.title}\n")
                        f.write(f"文号: {self._document.doc_number}\n")
                        f.write(f"有效性: {self._document.validity}\n")
                        f.write(f"来源: {self._document.url}\n")
                        f.write("\n" + "=" * 50 + "\n\n")
                        f.write(text)
                    
                    return True
            
            return False
            
        except Exception as e:
            print(f"提取正文失败: {e}")
            return False


class NormativePDFDownloadWorker(QThread):
    """规范性文件 PDF 下载工作线程
    
    直接下载网页中的 PDF 附件。
    优先使用浏览器模式获取页面内容（绕过 JS 反爬虫）。
    
    Feature: caac-regulation-download-enhancement
    Validates: Requirements 2.1, 2.2, 2.3, 2.4
    """
    
    # 信号
    finished = Signal(str)  # 下载完成，参数为文件路径
    error = Signal(str)  # 下载错误
    progress = Signal(int, int, str)  # 进度 (current, total, message)
    
    def __init__(
        self,
        session: requests.Session,
        document: RegulationDocument,
        save_path: str,
    ):
        """初始化
        
        Args:
            session: HTTP 会话
            document: 要下载的文档
            save_path: 保存路径
        """
        super().__init__()
        self._session = session
        self._document = document
        self._save_path = save_path
        self._should_stop = False
        self._browser_fetcher = None
        self._browser_fetcher_checked = False
    
    def _get_browser_fetcher(self):
        """获取浏览器获取器（懒加载）"""
        if self._browser_fetcher_checked:
            return self._browser_fetcher
        
        self._browser_fetcher_checked = True
        
        try:
            from screenshot_tool.services.browser_fetcher import BrowserFetcher
            if BrowserFetcher.is_available():
                self._browser_fetcher = BrowserFetcher(timeout=30)
                _debug_log("下载器：浏览器模式可用", "REGULATION")
            else:
                _debug_log("下载器：浏览器模式不可用", "REGULATION")
        except ImportError as e:
            _debug_log(f"下载器：browser_fetcher 导入失败: {e}", "REGULATION")
        except Exception as e:
            _debug_log(f"下载器：初始化浏览器获取器失败: {e}", "REGULATION")
        
        return self._browser_fetcher
    
    def _fetch_page_content(self, url: str) -> str:
        """获取页面内容
        
        优先使用浏览器模式（绕过 JS 反爬虫），失败时回退到 requests。
        
        Args:
            url: 页面 URL
            
        Returns:
            HTML 内容，失败时返回空字符串
        """
        _debug_log(f"下载器获取页面: {url}", "REGULATION")
        
        # 优先使用浏览器模式
        browser_fetcher = self._get_browser_fetcher()
        if browser_fetcher:
            try:
                result = browser_fetcher.fetch(url, use_cookies=False)
                if result.success and result.html:
                    _debug_log(f"下载器浏览器模式获取成功: {len(result.html)} 字符", "REGULATION")
                    return result.html
                else:
                    error_msg = result.error if hasattr(result, 'error') else "未知错误"
                    _debug_log(f"下载器浏览器模式获取失败: {error_msg}", "REGULATION")
            except Exception as e:
                _debug_log(f"下载器浏览器模式异常: {e}", "REGULATION")
        
        # 回退到 requests
        _debug_log("下载器回退到 requests 模式", "REGULATION")
        try:
            response = self._session.get(url, timeout=30)
            response.encoding = "utf-8"
            if response.status_code == 200:
                return response.text
            else:
                _debug_log(f"下载器 requests 获取失败: HTTP {response.status_code}", "REGULATION")
        except Exception as e:
            _debug_log(f"下载器 requests 获取异常: {e}", "REGULATION")
        
        return ""
    
    def stop(self):
        """请求停止"""
        self._should_stop = True
    
    def run(self):
        """执行下载
        
        1. 访问详情页（优先使用浏览器模式）
        2. 查找 PDF 附件链接
        3. 下载 PDF 文件
        4. 如果没有 PDF，尝试 DOC/DOCX
        5. 如果都没有，提取正文保存为 TXT
        """
        try:
            self.progress.emit(0, 100, "正在访问详情页...")
            
            # 访问详情页（优先使用浏览器模式）
            html_content = self._fetch_page_content(self._document.url)
            
            if not html_content:
                self.error.emit(f"访问详情页失败: 无法获取页面内容")
                return
            
            if self._should_stop:
                return
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            self.progress.emit(20, 100, "正在查找 PDF 附件...")
            
            # 查找 PDF 链接
            pdf_link = self._find_attachment_link(soup, ['.pdf'])
            
            if pdf_link:
                self.progress.emit(40, 100, "正在下载 PDF...")
                success = self._download_file(pdf_link, self._save_path)
                if success:
                    self.progress.emit(100, 100, "下载完成")
                    self.finished.emit(self._save_path)
                    return
            
            if self._should_stop:
                return
            
            # 尝试 DOC/DOCX
            self.progress.emit(50, 100, "正在查找 DOC/DOCX 附件...")
            doc_link = self._find_attachment_link(soup, ['.doc', '.docx'])
            
            if doc_link:
                self.progress.emit(60, 100, "正在下载 DOC/DOCX...")
                # 调整保存路径
                ext = '.docx' if '.docx' in doc_link.lower() else '.doc'
                doc_save_path = self._save_path.replace('.pdf', ext)
                success = self._download_file(doc_link, doc_save_path)
                if success:
                    self.progress.emit(100, 100, "下载完成")
                    self.finished.emit(doc_save_path)
                    return
            
            if self._should_stop:
                return
            
            # 提取正文保存为 TXT
            self.progress.emit(80, 100, "正在提取正文...")
            txt_path = self._save_path.replace(".pdf", ".txt")
            if self._extract_text_fallback(soup, txt_path):
                self.progress.emit(100, 100, "提取完成")
                self.finished.emit(txt_path)
            else:
                self.error.emit("未找到可下载的附件，正文提取也失败")
                
        except Exception as e:
            self.error.emit(f"下载失败: {str(e)}")
    
    def _find_attachment_link(self, soup: BeautifulSoup, extensions: list) -> Optional[str]:
        """查找附件链接
        
        Args:
            soup: BeautifulSoup 对象
            extensions: 要查找的文件扩展名列表
            
        Returns:
            附件链接或 None
        """
        # 构建正则表达式
        ext_pattern = '|'.join(re.escape(ext) for ext in extensions)
        pattern = re.compile(rf'({ext_pattern})$', re.I)
        
        # 模式1: 查找附件区域的链接
        attachment_elements = soup.find_all(string=re.compile(r'附件[：:]?', re.I))
        for attachment_text in attachment_elements:
            attachment_parent = attachment_text.parent
            if attachment_parent:
                parent_container = attachment_parent.parent if attachment_parent.parent else attachment_parent
                links = parent_container.find_all('a', href=pattern)
                if links:
                    return self._build_full_url(links[0].get('href'))
        
        # 模式2: 直接查找所有匹配的链接
        links = soup.find_all('a', href=pattern)
        if links:
            for link in links:
                href = link.get('href', '')
                if href and ('/XXGK/' in href or href.startswith('http')):
                    return self._build_full_url(href)
            return self._build_full_url(links[0].get('href'))
        
        return None
    
    def _build_full_url(self, link: str) -> str:
        """构建完整的 URL
        
        使用文档详情页 URL 的目录作为基准来解析相对路径。
        参考法规文件项目的 _build_pdf_url 实现。
        """
        if link.startswith('http'):
            return link
        
        # 处理相对路径前缀
        if link.startswith('./'):
            link = link[2:]
        elif link.startswith('../'):
            # 处理多级 ../ 
            doc_url = self._document.url
            while link.startswith('../'):
                link = link[3:]
                doc_url = '/'.join(doc_url.split('/')[:-1])
            doc_url_dir = '/'.join(doc_url.split('/')[:-1])
            return f"{doc_url_dir}/{link}"
        
        # 使用文档详情页 URL 的目录作为基准
        doc_url_dir = '/'.join(self._document.url.split('/')[:-1])
        return f"{doc_url_dir}/{link}"
    
    def _download_file(self, url: str, save_path: str) -> bool:
        """下载文件"""
        try:
            response = self._session.get(url, stream=True, timeout=60)
            
            if response.status_code != 200:
                return False
            
            total_size = int(response.headers.get("content-length", 0))
            downloaded = 0
            
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            with open(save_path, "wb") as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if self._should_stop:
                        f.close()
                        if os.path.exists(save_path):
                            os.remove(save_path)
                        return False
                    
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        
                        if total_size > 0:
                            percent = int(40 + (downloaded / total_size) * 50)
                            self.progress.emit(percent, 100, f"正在下载... {downloaded // 1024}KB")
            
            # 验证文件大小
            file_size = os.path.getsize(save_path)
            if file_size < 1024:
                os.remove(save_path)
                return False
            
            return True
            
        except Exception as e:
            print(f"下载文件失败: {e}")
            return False
    
    def _extract_text_fallback(self, soup: BeautifulSoup, save_path: str) -> bool:
        """提取正文作为备选方案"""
        try:
            content_div = soup.find("div", class_="TRS_Editor")
            if not content_div:
                content_div = soup.find("div", class_="content")
            if not content_div:
                content_div = soup.find("div", id="content")
            
            if content_div:
                text = content_div.get_text(separator="\n", strip=True)
                
                if text:
                    os.makedirs(os.path.dirname(save_path), exist_ok=True)
                    
                    with open(save_path, "w", encoding="utf-8") as f:
                        f.write(f"标题: {self._document.title}\n")
                        f.write(f"文号: {self._document.doc_number}\n")
                        f.write(f"有效性: {self._document.validity}\n")
                        f.write(f"来源: {self._document.url}\n")
                        f.write("\n" + "=" * 50 + "\n\n")
                        f.write(text)
                    
                    return True
            
            return False
            
        except Exception as e:
            print(f"提取正文失败: {e}")
            return False


class SmartDownloadWorker(QThread):
    """智能下载工作线程
    
    下载策略：
    1. 先检查详情页是否有 PDF 附件
    2. 如果有 PDF，直接下载 PDF
    3. 如果没有 PDF，提取内容生成 DOCX
    
    Feature: caac-regulation-download-enhancement
    """
    
    # 信号
    finished = Signal(str)  # 下载完成，参数为文件路径
    error = Signal(str)  # 下载错误
    progress = Signal(int, int, str)  # 进度 (current, total, message)
    
    def __init__(
        self,
        session: requests.Session,
        document: RegulationDocument,
        save_path: str,
    ):
        """初始化
        
        Args:
            session: HTTP 会话
            document: 要下载的文档
            save_path: 保存路径
        """
        super().__init__()
        self._session = session
        self._document = document
        self._save_path = save_path
        self._should_stop = False
        self._browser_fetcher = None
        self._browser_fetcher_checked = False
    
    def _get_browser_fetcher(self):
        """获取浏览器获取器（懒加载）"""
        if self._browser_fetcher_checked:
            return self._browser_fetcher
        
        self._browser_fetcher_checked = True
        
        try:
            from screenshot_tool.services.browser_fetcher import BrowserFetcher
            if BrowserFetcher.is_available():
                self._browser_fetcher = BrowserFetcher(timeout=30)
                _debug_log("智能下载器：浏览器模式可用", "REGULATION")
            else:
                _debug_log("智能下载器：浏览器模式不可用", "REGULATION")
        except ImportError as e:
            _debug_log(f"智能下载器：browser_fetcher 导入失败: {e}", "REGULATION")
        except Exception as e:
            _debug_log(f"智能下载器：初始化浏览器获取器失败: {e}", "REGULATION")
        
        return self._browser_fetcher
    
    def _fetch_page_content(self, url: str) -> str:
        """获取页面内容"""
        _debug_log(f"智能下载器获取页面: {url}", "REGULATION")
        
        browser_fetcher = self._get_browser_fetcher()
        if browser_fetcher:
            try:
                result = browser_fetcher.fetch(url, use_cookies=False)
                if result.success and result.html:
                    _debug_log(f"智能下载器浏览器模式获取成功: {len(result.html)} 字符", "REGULATION")
                    return result.html
            except Exception as e:
                _debug_log(f"智能下载器浏览器模式异常: {e}", "REGULATION")
        
        try:
            response = self._session.get(url, timeout=30)
            response.encoding = "utf-8"
            if response.status_code == 200:
                return response.text
        except Exception as e:
            _debug_log(f"智能下载器 requests 获取异常: {e}", "REGULATION")
        
        return ""
    
    def stop(self):
        """请求停止"""
        self._should_stop = True
    
    def run(self):
        """执行下载"""
        try:
            self.progress.emit(0, 100, "正在访问详情页...")
            
            html_content = self._fetch_page_content(self._document.url)
            
            if not html_content:
                self.error.emit("访问详情页失败: 无法获取页面内容")
                return
            
            if self._should_stop:
                return
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            self.progress.emit(20, 100, "正在查找 PDF 附件...")
            
            # 查找 PDF 链接
            pdf_link = self._find_attachment_link(soup, ['.pdf'])
            
            if pdf_link:
                self.progress.emit(40, 100, "正在下载 PDF...")
                if self._download_file(pdf_link, self._save_path):
                    self.progress.emit(100, 100, "下载完成")
                    self.finished.emit(self._save_path)
                    return
            
            if self._should_stop:
                return
            
            # 没有 PDF，生成 DOCX
            self.progress.emit(50, 100, "未找到 PDF，正在生成 DOCX...")
            
            docx_path = self._save_path
            if docx_path.endswith(".pdf"):
                docx_path = docx_path[:-4] + ".docx"
            elif not docx_path.endswith(".docx"):
                docx_path = docx_path + ".docx"
            
            content = self._extract_content(soup)
            
            if not content:
                self.error.emit("未能提取内容")
                return
            
            self.progress.emit(70, 100, "正在生成 DOCX 文件...")
            
            save_dir = os.path.dirname(docx_path)
            if save_dir:
                os.makedirs(save_dir, exist_ok=True)
            
            if self._convert_to_docx(content, docx_path):
                self.progress.emit(100, 100, "转换完成")
                self.finished.emit(docx_path)
            else:
                self.error.emit("DOCX 生成失败")
                
        except Exception as e:
            self.error.emit(f"下载失败: {str(e)}")
    
    def _find_attachment_link(self, soup: BeautifulSoup, extensions: list) -> Optional[str]:
        """查找附件链接"""
        ext_pattern = '|'.join(re.escape(ext) for ext in extensions)
        pattern = re.compile(rf'({ext_pattern})$', re.I)
        
        # 模式1: 查找附件区域的链接
        attachment_elements = soup.find_all(string=re.compile(r'附件[：:]?', re.I))
        for attachment_text in attachment_elements:
            attachment_parent = attachment_text.parent
            if attachment_parent:
                parent_container = attachment_parent.parent if attachment_parent.parent else attachment_parent
                links = parent_container.find_all('a', href=pattern)
                if links:
                    return self._build_full_url(links[0].get('href'))
        
        # 模式2: 直接查找所有匹配的链接
        links = soup.find_all('a', href=pattern)
        if links:
            for link in links:
                href = link.get('href', '')
                if href and ('/XXGK/' in href or href.startswith('http')):
                    return self._build_full_url(href)
            return self._build_full_url(links[0].get('href'))
        
        return None
    
    def _build_full_url(self, link: str) -> str:
        """构建完整的 URL"""
        if link.startswith('http'):
            return link
        
        if link.startswith('./'):
            link = link[2:]
        elif link.startswith('../'):
            doc_url = self._document.url
            while link.startswith('../'):
                link = link[3:]
                doc_url = '/'.join(doc_url.split('/')[:-1])
            doc_url_dir = '/'.join(doc_url.split('/')[:-1])
            return f"{doc_url_dir}/{link}"
        
        doc_url_dir = '/'.join(self._document.url.split('/')[:-1])
        return f"{doc_url_dir}/{link}"
    
    def _download_file(self, url: str, save_path: str) -> bool:
        """下载文件"""
        try:
            response = self._session.get(url, stream=True, timeout=60)
            
            if response.status_code != 200:
                return False
            
            total_size = int(response.headers.get("content-length", 0))
            downloaded = 0
            
            save_dir = os.path.dirname(save_path)
            if save_dir:
                os.makedirs(save_dir, exist_ok=True)
            
            with open(save_path, "wb") as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if self._should_stop:
                        f.close()
                        if os.path.exists(save_path):
                            os.remove(save_path)
                        return False
                    
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        
                        if total_size > 0:
                            percent = int(40 + (downloaded / total_size) * 50)
                            self.progress.emit(percent, 100, f"正在下载... {downloaded // 1024}KB")
            
            file_size = os.path.getsize(save_path)
            if file_size < 1024:
                os.remove(save_path)
                return False
            
            return True
            
        except Exception as e:
            _debug_log(f"下载文件失败: {e}", "REGULATION")
            return False
    
    def _extract_content(self, soup: BeautifulSoup) -> str:
        """提取内容"""
        content_div = soup.find("div", class_="TRS_Editor")
        if not content_div:
            content_div = soup.find("div", class_="content")
        if not content_div:
            content_div = soup.find("div", id="content")
        if not content_div:
            content_div = soup.find("div", class_="article-content")
        
        if content_div:
            return str(content_div)
        
        return ""
    
    def _convert_to_docx(self, html_content: str, save_path: str) -> bool:
        """转换为 DOCX 格式"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            _debug_log("python-docx 未安装", "REGULATION")
            return False
        
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            doc = Document()
            
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            
            title = doc.add_heading(self._document.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"文号: {self._document.doc_number or '无'}").bold = True
            meta_para.add_run(f"    有效性: {self._document.validity or '未知'}").bold = True
            meta_para.add_run(f"\n来源: {self._document.url}")
            
            doc.add_paragraph()
            
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'ul', 'ol']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    text = element.get_text(strip=True)
                    if text:
                        doc.add_heading(text, level=min(level, 9))
                elif element.name == 'p':
                    text = element.get_text(strip=True)
                    if text:
                        if re.match(r'^第[一二三四五六七八九十百]+章', text):
                            heading = doc.add_heading(text, level=1)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif re.match(r'^第[一二三四五六七八九十百]+节', text):
                            heading = doc.add_heading(text, level=2)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif re.match(r'^第[一二三四五六七八九十百\d]+条', text):
                            para = doc.add_paragraph()
                            para.add_run(text).bold = True
                        else:
                            para = doc.add_paragraph(text)
                            para.paragraph_format.first_line_indent = Inches(0.3)
                elif element.name == 'table':
                    text = element.get_text(separator="\n", strip=True)
                    if text:
                        para = doc.add_paragraph(text)
                        para.style = 'Quote'
                elif element.name in ('ul', 'ol'):
                    for li in element.find_all('li'):
                        text = li.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text, style='List Bullet')
            
            doc.save(save_path)
            return True
            
        except Exception as e:
            _debug_log(f"DOCX 转换失败: {e}", "REGULATION")
            return False


class RegulationDocxDownloadWorker(QThread):
    """CCAR 规章 DOCX 下载工作线程
    
    提取内容直接转换为 DOCX 格式。
    优先使用浏览器模式获取页面内容（绕过 JS 反爬虫）。
    
    Feature: caac-regulation-download-enhancement
    Validates: Requirements 3.1, 3.2, 3.3, 3.4, 3.5
    """
    
    # 信号
    finished = Signal(str)  # 下载完成，参数为文件路径
    error = Signal(str)  # 下载错误
    progress = Signal(int, int, str)  # 进度 (current, total, message)
    
    def __init__(
        self,
        session: requests.Session,
        document: RegulationDocument,
        save_path: str,
    ):
        """初始化
        
        Args:
            session: HTTP 会话
            document: 要下载的文档
            save_path: 保存路径（会自动改为 .docx 后缀）
        """
        super().__init__()
        self._session = session
        self._document = document
        # 确保使用 .docx 后缀
        if save_path.endswith(".pdf"):
            save_path = save_path[:-4] + ".docx"
        elif not save_path.endswith(".docx"):
            save_path = save_path + ".docx"
        self._save_path = save_path
        self._should_stop = False
        self._browser_fetcher = None
        self._browser_fetcher_checked = False
    
    def _get_browser_fetcher(self):
        """获取浏览器获取器（懒加载）"""
        if self._browser_fetcher_checked:
            return self._browser_fetcher
        
        self._browser_fetcher_checked = True
        
        try:
            from screenshot_tool.services.browser_fetcher import BrowserFetcher
            if BrowserFetcher.is_available():
                self._browser_fetcher = BrowserFetcher(timeout=30)
                _debug_log("CCAR下载器：浏览器模式可用", "REGULATION")
            else:
                _debug_log("CCAR下载器：浏览器模式不可用", "REGULATION")
        except ImportError as e:
            _debug_log(f"CCAR下载器：browser_fetcher 导入失败: {e}", "REGULATION")
        except Exception as e:
            _debug_log(f"CCAR下载器：初始化浏览器获取器失败: {e}", "REGULATION")
        
        return self._browser_fetcher
    
    def _fetch_page_content(self, url: str) -> str:
        """获取页面内容
        
        优先使用浏览器模式（绕过 JS 反爬虫），失败时回退到 requests。
        """
        _debug_log(f"CCAR下载器获取页面: {url}", "REGULATION")
        
        browser_fetcher = self._get_browser_fetcher()
        if browser_fetcher:
            try:
                result = browser_fetcher.fetch(url, use_cookies=False)
                if result.success and result.html:
                    _debug_log(f"CCAR下载器浏览器模式获取成功: {len(result.html)} 字符", "REGULATION")
                    return result.html
                else:
                    error_msg = result.error if hasattr(result, 'error') else "未知错误"
                    _debug_log(f"CCAR下载器浏览器模式获取失败: {error_msg}", "REGULATION")
            except Exception as e:
                _debug_log(f"CCAR下载器浏览器模式异常: {e}", "REGULATION")
        
        _debug_log("CCAR下载器回退到 requests 模式", "REGULATION")
        try:
            response = self._session.get(url, timeout=30)
            response.encoding = "utf-8"
            if response.status_code == 200:
                return response.text
            else:
                _debug_log(f"CCAR下载器 requests 获取失败: HTTP {response.status_code}", "REGULATION")
        except Exception as e:
            _debug_log(f"CCAR下载器 requests 获取异常: {e}", "REGULATION")
        
        return ""
    
    def stop(self):
        """请求停止"""
        self._should_stop = True
    
    def run(self):
        """执行下载
        
        1. 访问详情页（优先使用浏览器模式）
        2. 提取规章内容
        3. 转换为 DOCX 格式
        """
        try:
            self.progress.emit(0, 100, "正在访问详情页...")
            
            # 访问详情页（优先使用浏览器模式）
            html_content = self._fetch_page_content(self._document.url)
            
            if not html_content:
                self.error.emit(f"访问详情页失败: 无法获取页面内容")
                return
            
            if self._should_stop:
                return
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            self.progress.emit(30, 100, "正在提取规章内容...")
            
            # 提取规章内容
            content = self._extract_regulation_content(soup)
            
            if not content:
                self.error.emit("未能提取规章内容")
                return
            
            if self._should_stop:
                return
            
            self.progress.emit(60, 100, "正在生成 DOCX 文件...")
            
            # 转换为 DOCX
            save_dir = os.path.dirname(self._save_path)
            if save_dir:
                os.makedirs(save_dir, exist_ok=True)
            
            if self._convert_to_docx(content, self._save_path):
                self.progress.emit(100, 100, "转换完成")
                self.finished.emit(self._save_path)
            else:
                self.error.emit("DOCX 生成失败")
                
        except Exception as e:
            self.error.emit(f"下载失败: {str(e)}")
    
    def _extract_regulation_content(self, soup: BeautifulSoup) -> str:
        """提取规章内容
        
        Args:
            soup: BeautifulSoup 对象
            
        Returns:
            提取的内容文本
        """
        # 查找正文内容
        content_div = soup.find("div", class_="TRS_Editor")
        if not content_div:
            content_div = soup.find("div", class_="content")
        if not content_div:
            content_div = soup.find("div", id="content")
        if not content_div:
            content_div = soup.find("div", class_="article-content")
        
        if content_div:
            # 获取 HTML 内容以保留结构
            return str(content_div)
        
        return ""
    
    def _convert_to_docx(self, html_content: str, save_path: str) -> bool:
        """转换为 DOCX 格式
        
        Args:
            html_content: HTML 内容
            save_path: 保存路径
            
        Returns:
            是否成功
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            _debug_log("python-docx 未安装，尝试安装...", "REGULATION")
            try:
                import subprocess
                subprocess.run(
                    ["pip", "install", "python-docx"],
                    capture_output=True,
                    timeout=60
                )
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except Exception as e:
                _debug_log(f"安装 python-docx 失败: {e}", "REGULATION")
                return False
        
        try:
            # 解析 HTML
            soup = BeautifulSoup(html_content, "html.parser")
            
            # 创建文档
            doc = Document()
            
            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            
            # 添加标题
            title = doc.add_heading(self._document.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元数据
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"文号: {self._document.doc_number or '无'}").bold = True
            meta_para.add_run(f"    有效性: {self._document.validity or '未知'}").bold = True
            meta_para.add_run(f"\n来源: {self._document.url}")
            
            doc.add_paragraph()  # 空行
            
            # 转换内容
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'ul', 'ol']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    text = element.get_text(strip=True)
                    if text:
                        doc.add_heading(text, level=min(level, 9))
                elif element.name == 'p':
                    text = element.get_text(strip=True)
                    if text:
                        # 检测是否是章节标题
                        if re.match(r'^第[一二三四五六七八九十百]+章', text):
                            heading = doc.add_heading(text, level=1)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif re.match(r'^第[一二三四五六七八九十百]+节', text):
                            heading = doc.add_heading(text, level=2)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif re.match(r'^第[一二三四五六七八九十百\d]+条', text):
                            para = doc.add_paragraph()
                            para.add_run(text).bold = True
                        else:
                            # 普通段落，首行缩进
                            para = doc.add_paragraph(text)
                            para.paragraph_format.first_line_indent = Inches(0.3)
                elif element.name == 'table':
                    # 简单处理表格为文本
                    text = element.get_text(separator="\n", strip=True)
                    if text:
                        para = doc.add_paragraph(text)
                        para.style = 'Quote'
                elif element.name in ('ul', 'ol'):
                    for li in element.find_all('li'):
                        text = li.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text, style='List Bullet')
            
            # 保存文档
            doc.save(save_path)
            return True
            
        except Exception as e:
            _debug_log(f"DOCX 转换失败: {e}", "REGULATION")
            return False


# 保留旧类名作为别名，保持向后兼容
RegulationMarkdownDownloadWorker = RegulationDocxDownloadWorker


class BatchDownloadManager(QObject):
    """批量下载管理器
    
    Feature: caac-regulation-download-enhancement
    Validates: Requirements 7.3, 7.4, 7.5
    """
    
    # 信号
    progressChanged = Signal(int, int, str)  # (current, total, message)
    documentCompleted = Signal(str, str)  # (doc_title, file_path)
    documentFailed = Signal(str, str)  # (doc_title, error_msg)
    allCompleted = Signal(int, int)  # (success_count, fail_count)
    
    def __init__(self, service: RegulationService):
        """初始化
        
        Args:
            service: 规章服务实例
        """
        super().__init__()
        self._service = service
        self._documents: list[RegulationDocument] = []
        self._current_index = 0
        self._success_count = 0
        self._fail_count = 0
        self._is_running = False
        self._should_stop = False
        self._current_worker: Optional[QThread] = None
    
    def start_batch_download(self, documents: list[RegulationDocument]) -> None:
        """开始批量下载
        
        Args:
            documents: 要下载的文档列表
        """
        if self._is_running:
            return
        
        self._documents = documents
        self._current_index = 0
        self._success_count = 0
        self._fail_count = 0
        self._is_running = True
        self._should_stop = False
        
        self._download_next()
    
    def cancel(self) -> None:
        """取消下载"""
        self._should_stop = True
        if self._current_worker and self._current_worker.isRunning():
            self._disconnect_worker_signals()
            self._current_worker.stop()
            self._current_worker.wait(1000)  # 等待线程结束
        self._current_worker = None
    
    def _download_next(self) -> None:
        """下载下一个文档"""
        if self._should_stop or self._current_index >= len(self._documents):
            self._finish()
            return
        
        doc = self._documents[self._current_index]
        total = len(self._documents)
        
        self.progressChanged.emit(
            self._current_index + 1,
            total,
            f"正在下载 {self._current_index + 1}/{total}: {doc.title[:30]}..."
        )
        
        # 获取保存路径
        save_path = get_save_path(doc, self._service.save_path)
        
        # 根据文档类型选择下载策略
        if doc.doc_type == "normative":
            self._current_worker = NormativePDFDownloadWorker(
                self._service._session, doc, save_path
            )
        else:
            self._current_worker = RegulationMarkdownDownloadWorker(
                self._service._session, doc, save_path
            )
        
        self._current_worker.finished.connect(self._on_document_finished)
        self._current_worker.error.connect(self._on_document_error)
        self._current_worker.start()
    
    def _on_document_finished(self, file_path: str) -> None:
        """单个文档下载完成"""
        # 断开信号连接，避免重复触发
        self._disconnect_worker_signals()
        
        doc = self._documents[self._current_index]
        self._success_count += 1
        self.documentCompleted.emit(doc.title, file_path)
        
        self._current_index += 1
        self._download_next()
    
    def _on_document_error(self, error_msg: str) -> None:
        """单个文档下载失败"""
        # 断开信号连接，避免重复触发
        self._disconnect_worker_signals()
        
        doc = self._documents[self._current_index]
        self._fail_count += 1
        self.documentFailed.emit(doc.title, error_msg)
        
        # 继续下载下一个
        self._current_index += 1
        self._download_next()
    
    def _disconnect_worker_signals(self) -> None:
        """断开当前 worker 的信号连接"""
        if self._current_worker:
            try:
                self._current_worker.finished.disconnect(self._on_document_finished)
                self._current_worker.error.disconnect(self._on_document_error)
            except (RuntimeError, TypeError):
                # 信号可能未连接或已断开
                pass
    
    def _finish(self) -> None:
        """完成批量下载"""
        self._is_running = False
        self.allCompleted.emit(self._success_count, self._fail_count)
